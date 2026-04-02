from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style.paragraph_format.line_spacing = 1.5

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.name = '맑은 고딕'
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.name = '맑은 고딕'
    return p

def para(text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.font.size = Pt(10)
    return p

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = '맑은 고딕'
                r.font.size = Pt(9)
        set_cell_shading(hdr[i], '333333')
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = t.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = '맑은 고딕'
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return t

def org_chart(lines):
    """monospace org chart"""
    p = doc.add_paragraph()
    for line in lines:
        run = p.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    doc.add_paragraph()

def spacer():
    doc.add_paragraph()

# ═══════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════
spacer(); spacer(); spacer()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('씨앤씨인터내셔널')
r.bold = True; r.font.size = Pt(16); r.font.name = '맑은 고딕'
r.font.color.rgb = RGBColor(0x66,0x66,0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('조직구조 개편 제안')
r.bold = True; r.font.size = Pt(24); r.font.name = '맑은 고딕'

spacer()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('2026년 4월  |  경영기획본부')
r.font.size = Pt(11); r.font.name = '맑은 고딕'
r.font.color.rgb = RGBColor(0x88,0x88,0x88)

doc.add_page_break()

# ═══════════════════════════════════════════
# I. 우리가 가야 할 곳
# ═══════════════════════════════════════════
h1('I. 우리가 가야 할 곳 — Seamless, 그리고 청주')

h2('경쟁사들은 경영목표를 조직구조에 심었다')

para('코스맥스·콜마·인터코스가 1조를 넘긴 공통점은 하나다.')
para('"고객 요청 → PD → 생산 → 납품"의 흐름을 조직구조에 내재화해서, 전략이 실행으로 직결되는 체계를 만들었다는 것.', bold=True)
para('전략은 따로 있고 실행은 따로 움직이는 구조로는, 아무리 수주를 늘려도 속도가 나지 않는다.')
spacer()

table(
    ['회사', '매출', '핵심 전환', '성과'],
    [
        ['코스맥스', '2조 1,000억', 'SCM 독립 C-Level화', '리드타임 22% 단축'],
        ['콜마코리아', '1조 8,000억', 'COO + CBO 도입', '영업이익 +85.8%'],
        ['인터코스', '1조 2,000억', 'Chairman/CEO 이원체제', '아시아 +24.3%'],
        ['씨앤씨', '5,243억', '—', '영업이익률 9.2%'],
    ]
)

h2('데드라인: 청주공장 2027년 하반기, 지금부터 15~18개월')

para('청주공장이 열리는 순간, 조달·생산계획·영업관리가 동시에 폭발적으로 늘어난다.')
para('지금의 구조로는 그 규모를 소화할 수 없다.')
para('조직이 공장보다 먼저 준비되어야 한다.', bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════
# II. 오퍼레이션만
# ═══════════════════════════════════════════
h1('II. 씨앤씨는 그간 오퍼레이션 기능만 가지고 있을 뿐, 전략과 효율화는 미비하다')

h2('1. 전략 기능이 없다')

para('현재 씨앤씨에는 FP&A, 전략기획, 사업개발 기능이 사실상 부재하다.')
bullet('재경팀은 결산·세무·자금 집행에 집중 → 미래 수익성 시뮬레이션 불가')
bullet('경영기획본부는 경영정보·인사·총무 중심 → 성장 전략 수립 주체 없음')
spacer()
para('"어디서 얼마를 벌어야 하는가"를 설계하는 기능이 없다.', bold=True)
para('청주공장의 용량을 채울 수주 전략, 채널·고객 포트폴리오 설계, 수익성 관리. 이 판단을 내릴 구조 자체가 지금 없다.')
spacer()

h2('2. 효율화 기능이 없다 — 그마저도 작동하지 않는다')

para('① 인사: 1,257명 조직을 채용담당 1인이 관할', bold=True)
para('전사 인원 1,257명, 월 인건비 47.1억. 청주공장 가동 시 대규모 채용이 불가피하나, 현재 인사팀은 채용 1인 + 노무·급여 중심으로 운영된다. 전략적 인재 확보 기능이 없다.')
spacer()

para('② SCM: 매출 66.8%(1,784억)에 영향을 주는 기능이 생산본부 하위에 묻혀 있다', bold=True)
para('전략구매·생산운영·영업관리는 매출과 수익성에 직결된 핵심 기능이다. 그러나 현재 이 기능은 생산본부 하위의 생산기획부로 편제되어 있어, 속도와 우선순위가 생산 논리에 종속된다. 납기 이슈가 생겨도, 구매 전략을 바꿔야 해도, SCM이 독립적으로 움직일 수 없다.')
spacer()

h2('3. 공동경영 구조가 경영진 레벨까지 내려오고 있다')

para('현재 씨앤씨의 공동경영 논리는 이사회를 넘어 경영진 포지션 배분에까지 영향을 준다.')
bullet('기능이 아닌 서열 기반으로 직책이 배분됨')
bullet('신임 경영진(전무)과 기존 경영진(부사장) 간 권한 충돌 구조 내재')
bullet('C-Level로 전환해도 직함만 바뀌고 권한은 서열에 종속될 위험')
spacer()
para('경쟁사 4사 모두 공통점은 하나다. 오너는 이사회에 머물고, 경영진은 기능·능력 기반으로 전권을 행사한다. 지배구조(이사회)와 경영진을 분리하는 것이 선결 조건이다.', bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════
# III. 현행 조직구조
# ═══════════════════════════════════════════
h1('III. 현행 조직구조 — 실제로 어떻게 생겼나')

para('(2026년 인사 자료 기준, 전사 1,257명)')
spacer()

para('현행 조직도', bold=True, size=12)

org_chart([
    '                              CEO',
    '                               │',
    '                ┌──────────────┴──────────────┐',
    '                │                             │',
    '           부사장 (1)                     부사장 (2)',
    '         경영·생산 관할                  개발·영업 관할',
    '                │                             │',
    '         ┌──────┴──────┐          ┌───────────┼───────────┐',
    '    경영기획본부    생산본부    제품개발본부    OD본부      연구소',
    '    경영정보           │       OBM            제품기획    MU / SC',
    '    인사               │       KPD 1~4',
    '    재경               │       GPD 1~3',
    '    재무기획           │',
    '    총무               │',
    '               ┌───────┼───────┐',
    '          생산기획부  퍼플공장  그린공장  3공장',
    '          전략구매',
    '          생산운영',
    '          영업관리',
])

h2('구조의 본질적 문제')

para('① 부사장(1)에 조직이 과집중되어 있다', bold=True)
para('전사에서 가장 많은 인원과 매출 영향력을 가진 기능 — 생산(980명)과 SCM(매출 66.8% 영향) — 이 모두 부사장(1) 한 명 아래 집중되어 있다. 여기에 경영기획본부(인사·재무·총무)까지 관할한다. CEO가 전략을 판단하려 해도, 운영 정보의 대부분이 이 한 명의 라인을 통해서만 올라오는 구조다.')
spacer()

para('② 부사장은 SCM 전문가가 아니다', bold=True)
para('현재 SCM 기능(전략구매·생산운영·영업관리)을 관할하는 부사장은 생산·품질 배경의 실행 중심 인력이다. 전략구매, S&OP 설계, 납기 최적화는 별개의 전문 영역임에도, 생산 논리가 SCM 판단을 지배하고 있다. 코스맥스가 SCM에 독립 C-Level을 둔 이유가 정확히 이것이다.')
spacer()

para('③ 영업관리는 전략이 아닌 실행 위주 인력으로 구성되어 있다', bold=True)
para('영업관리팀이 생산기획부 하위에 편제되어 있고, 구성 인력은 수주 처리·납기 조율 등 실행 중심이다. 시장 전체를 보고 고객 포트폴리오를 설계하거나, 채널별 수익성을 분석해 수주 전략을 짜는 기능이 없다. 영업이 실행만 하고 전략은 없는 구조다.')
spacer()

h2('현행 구조 문제 요약')
table(
    ['문제', '현황', '영향'],
    [
        ['부사장(1) 관할 과집중', '경영기획+생산 980명+SCM 전부', 'CEO 정보 접근 제한, 판단 왜곡'],
        ['SCM 전문성 부재', '생산 배경 임원이 SCM 관할', '구매·납기 전략이 생산 논리에 종속'],
        ['영업 전략 기능 없음', '수주 실행 중심 인력', '고객 포트폴리오·채널 전략 공백'],
        ['전략기획 주체 없음', '경영기획본부 = 관리 기능 묶음', '청주 이후 성장 설계 불가'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════
# IV. 1안 vs 2안
# ═══════════════════════════════════════════
h1('IV. 그렇다면 어떤 구조로 가야 하는가 — 1안 vs 2안')

para('씨앤씨가 선택할 수 있는 경로는 두 가지다.')
para('어떤 안을 택하든, 성과에 대한 책임은 반드시 특정 C-Level에 귀속되어야 한다.', bold=True)
spacer()

# ─── 1안 ───
h2('1안: 코스맥스형 — Chief SCM 독립')

org_chart([
    '                              CEO',
    '                               │',
    '          ┌────────────────────┼────────────────────┐',
    '          │          │         │          │',
    '        CFO         CBO    Chief SCM    생산총괄',
    '     (재무·인사)  (전략·사업)  (독립 C-Level)  (품질·기술)',
    '          │          │         │',
    '       FP&A팀     전략기획팀   전략구매팀',
    '       재경팀      OD본부      생산운영팀 (S&OP)',
    '       인사팀      제품개발본부  영업관리팀',
    '                               │',
    '                             공장장',
    '                    (퍼플·그린·3공장·청주)',
])

para('성과 책임: Chief SCM이 납기율·조달비용·재고회전율 전부를 단독 책임진다. 공장장이 SCM 직보이므로 생산 일정이 SCM 논리로 움직인다.', bold=True)
spacer()

table(
    ['장점', '단점'],
    [
        ['SCM이 생산 우선순위 직접 통제 → 납기·속도 최적화', 'Chief SCM 1인에 권한 집중 → 인물 리스크 큼'],
        ['청주공장 추가 시 공장장 즉시 편입', '생산총괄과 공장장 지휘 권한 충돌 가능'],
        ['성과 KPI 단일화, 책임 추적 명확', '외부 영입 필수 → 적응 기간 중 공백 위험'],
        ['코스맥스 실증: 리드타임 22% 단축', '내부 발탁으로 채울 적임자 현재 없음'],
    ]
)

# ─── 2안 ───
h2('2안: 콜마형 — COO 산하 SCM+생산 통합')

org_chart([
    '                              CEO',
    '                               │',
    '               ┌───────────────┼───────────────┐',
    '               │               │               │',
    '             CFO              CBO             COO',
    '          (재무·인사)       (전략·사업)   (SCM+생산 통합)',
    '               │               │               │',
    '            FP&A팀          전략기획팀     ┌────┴────┐',
    '            재경팀           OD본부      SCM      생산총괄',
    '            인사팀           제품개발본부  전략구매   공장장(전공장)',
    '                                        생산운영   품질·기술팀',
    '                                        영업관리',
])

para('성과 책임: COO가 매출원가·납기율·생산효율 전부를 책임진다. SCM과 생산이 같은 지휘 라인 아래 있어 내부 마찰을 COO가 중재한다.', bold=True)
spacer()

table(
    ['장점', '단점'],
    [
        ['SCM-생산 충돌을 COO가 중재 → 현장 안정성', 'COO가 생산 논리에 기울면 SCM 다시 종속'],
        ['내부 발탁으로 전환 충격 최소화', '청주 가동 후 공장 4개 + SCM 전부 관할 → 과부하'],
        ['콜마 실증: 영업이익 +85.8%', 'KPI가 COO에 혼재 → 성과 추적 복잡'],
        ['기존 경영진과의 마찰 상대적으로 적음', 'COO 포지션 없으면 구조 자체가 성립 안 됨'],
    ]
)

# ─── 비교 ───
h2('핵심 비교')
table(
    ['항목', '1안 (코스맥스형)', '2안 (콜마형)'],
    [
        ['성과 책임 소재', 'Chief SCM 단독, 명확', 'COO 단독, 범위 과대'],
        ['납기 KPI 추적', 'SCM에 직결', 'COO KPI에 혼재'],
        ['생산-SCM 충돌', '구조적으로 발생 가능', 'COO가 중재'],
        ['청주공장 대응', '공장장 바로 편입', 'COO 부하 급증'],
        ['전환 난이도', '외부 영입 필수', '내부 발탁 가능'],
        ['장기 확장성', '높음', 'COO 과부하로 한계 도달'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════
# V. 권고안
# ═══════════════════════════════════════════
h1('V. 권고안 — 2안으로 시작, 1안으로 전환')

para('두 안은 택일의 문제가 아니라 시점의 문제다.', bold=True, size=11)
spacer()

para('Phase 1 (현재 ~ 청주 가동 전): 2안 — COO 체계로 전환', bold=True)
bullet('외부 영입 리스크 없이 내부 안정적 전환')
bullet('COO가 SCM+생산을 통합 관할하며 청주 준비 주도')
bullet('이 시점에 Chief SCM 외부 후보 탐색 병행')
spacer()

para('Phase 2 (청주 가동 시점): 1안 — Chief SCM 독립', bold=True)
bullet('공장 4개 체계에서 COO 단독 관할은 구조적 과부하')
bullet('Chief SCM을 독립시켜 납기·조달·S&OP 전담')
bullet('COO는 생산품질·기술 총괄로 역할 재정의 또는 역할 종료')
spacer()

table(
    ['시점', '구조', '핵심 과제'],
    [
        ['지금', '현행 구조', 'C-Level 기능 정의 착수'],
        ['→ 전환기', '2안 (COO 체계)', '내부 안정화 + Chief SCM 외부 탐색'],
        ['→ 청주 가동 후', '1안 (Chief SCM 독립)', '성과 책임 구조 완성'],
    ]
)

para('핵심 전제: COO 선임 직후부터 Chief SCM 후보를 탐색해야 한다. 청주 가동 6개월 전에 Chief SCM이 온보딩되어 있어야 전환이 가능하다.', bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════
# VI. 후속 조치
# ═══════════════════════════════════════════
h1('VI. 후속 조치')
table(
    ['단계', '시점', '내용', '책임'],
    [
        ['1', '즉시', 'CFO·CBO 기능 정의, FP&A·전략기획팀 설계', 'CEO 승인'],
        ['2', '1개월 내', 'COO 후보 선정 (내부 발탁 우선)', '경영진'],
        ['3', '3개월 내', 'SCM KPI 체계 설계 (납기율·조달비용·재고)', 'COO'],
        ['4', '6개월 내', 'Chief SCM 외부 후보 탐색 개시', 'CFO + 헤드헌터'],
        ['5', '청주 가동 6개월 전', 'Chief SCM 온보딩 완료', 'CEO'],
        ['6', '청주 가동 시', '1안 전환 완료, 공장장 직보 체계 가동', 'Chief SCM'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 별첨 1
# ═══════════════════════════════════════════
h1('[별첨 1] 경쟁사 경영진 지배구조 비교')
table(
    ['', '코스맥스', '콜마', '인터코스', '코스메카', '씨앤씨'],
    [
        ['경영진 구조', '오너+전문경영인\n+부문장 3단', '지주(오너)+사업사\n(전문경영인)', '회장(오너)+CEO\n(전문경영인)', '부회장(내부)\n+사장(외부)', '공동경영 논리가\n경영진까지'],
        ['의사결정 속도', '부문 자율 전결', '지주-사업사 병렬', 'CEO 단일 집행', '병렬 독립 전결', '정렬 비용 발생'],
        ['SCM 구조', '구매·생산·물류\n독립 본부', 'COO 산하 통합', '수요 기반\n공장 배분', 'APS 기반\n생산계획', '생산기획부\n아래 혼재'],
        ['납기 효율', '리드타임 22% 단축', '기술영업으로\n개발속도 향상', '16개 공장 표준화', 'APS+실시간 관제', '측정 기능 없음'],
        ['영업이익 성장', '+51.6%', '+85.8%', '+지속 성장', '+58.4%', '9.2%\n(구조 개선 여지)'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 별첨 2
# ═══════════════════════════════════════════
h1('[별첨 2] Chief SCM 후보자 리서치')

para('씨앤씨의 Chief SCM은 S&OP·전략구매·생산계획·영업관리를 통합하고 청주공장 가동에 맞춰 전사 공급망을 설계할 수 있는 인물이어야 한다.')
spacer()

table(
    ['후보', '배경', '강점', '고려사항'],
    [
        ['김유태', '코스맥스 SCM본부 출신', 'S&OP 설계 직접 경험, ODM 도메인', '현직 이탈 조건 확인 필요'],
        ['박진호', 'LG생활건강 구매본부', '전략구매·원가협상, 대형 조직 경험', 'ODM 특수성 적응 기간 필요'],
        ['이동순', '아모레퍼시픽 생산기획', '생산계획·납기 관리 전문', 'SCM 전체 범위 경험 확인 필요'],
        ['정창욱', '코스메카 공급망', '중소 ODM SCM 실무 경험', '규모 확장 리더십 검증 필요'],
        ['류재민', '글로벌 3PL 출신', '물류·조달 통합 시각', '화장품 ODM 도메인 러닝커브'],
    ]
)

para('영입 전략: 코스맥스·코스메카 SCM 출신 중 S&OP 설계 직접 경험자 1순위. COO 선임 직후 헤드헌터 동시 접촉 시작.', bold=True)

spacer(); spacer()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('본 문서는 내부 경영진 논의용이며 외부 배포를 금합니다.')
r.font.color.rgb = RGBColor(0x99,0x99,0x99)
r.font.size = Pt(9); r.font.name = '맑은 고딕'

path = 'C:/Users/user/Desktop/씨앤씨_조직구조_개편제안_최종.docx'
doc.save(path)
print(f'저장 완료: {path}')
